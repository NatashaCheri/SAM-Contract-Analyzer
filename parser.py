"""
Document parser for the SAM Vendor Contract Analyzer.

Responsibilities:
- Accept an uploaded PDF or DOCX file (as bytes, from Streamlit's file_uploader)
- Extract text page-by-page (PDF) or in reading order (DOCX)
- Fall back to OCR automatically for pages that have little/no extractable
  text (i.e. scanned/image-only pages)
- Return a structured result that preserves page numbers, so later on the
  LLM extraction step can cite "page N" for every field it pulls out

No data is written to permanent storage. Everything lives in a per-session
temp directory (see SessionWorkspace below) that the caller is responsible
for cleaning up at the end of the session.
"""

from __future__ import annotations

import csv
import io
import shutil
import tempfile
import uuid
from dataclasses import dataclass, field
from pathlib import Path

import pymupdf  # PyMuPDF, imported as pymupdf (fitz is the deprecated alias)
import pytesseract
from docx import Document as DocxDocument
from PIL import Image, ImageOps

# ---------------------------------------------------------------------------
# Config
# ---------------------------------------------------------------------------

# If a PDF page yields fewer than this many characters of extractable text,
# we treat it as a scanned/image page and OCR it instead.
MIN_CHARS_FOR_NATIVE_TEXT = 20

# A page can have SOME extractable text and still be mostly a scan -- a very
# common pattern in real contract packs, where a scanned order form carries a
# born-digital header, footer or Bates stamp. Such a page clears the 20-char
# bar above while the actual licence table on it stays invisible. If a page
# has less than this much text but contains a large image, OCR it as well and
# keep whichever result is richer.
SPARSE_TEXT_THRESHOLD = 350
LARGE_IMAGE_AREA_RATIO = 0.35  # image covering >35% of the page

# DPI to render PDF pages at before handing them to Tesseract. Higher = more
# accurate OCR but slower. 300 is a good default for contract scans.
OCR_RENDER_DPI = 300

# Tesseract page segmentation mode 6 ("assume a uniform block of text") beats
# the default on order forms and licence schedules, where the default mode
# tends to shred table rows into disconnected fragments.
OCR_CONFIG = "--psm 6"

IMAGE_SUFFIXES = (".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".webp")
SUPPORTED_SUFFIXES = (
    ".pdf", ".docx", ".xlsx", ".xlsm", ".csv", ".txt",
) + IMAGE_SUFFIXES


# ---------------------------------------------------------------------------
# Data model
# ---------------------------------------------------------------------------

@dataclass
class PageResult:
    page_number: int  # 1-indexed
    text: str
    source: str  # "native" or "ocr"


@dataclass
class ParsedDocument:
    filename: str
    file_type: str  # "pdf" or "docx"
    pages: list[PageResult] = field(default_factory=list)

    @property
    def full_text(self) -> str:
        """Concatenated text with page markers, ready to feed to an LLM."""
        parts = []
        for p in self.pages:
            parts.append(f"\n\n--- Page {p.page_number} ---\n{p.text}")
        return "".join(parts).strip()

    @property
    def ocr_page_count(self) -> int:
        return sum(1 for p in self.pages if p.source == "ocr")

    @property
    def char_count(self) -> int:
        return sum(len(p.text) for p in self.pages)


# ---------------------------------------------------------------------------
# Session workspace (temp-only storage, no persistence)
# ---------------------------------------------------------------------------

class SessionWorkspace:
    """
    A throwaway temp directory scoped to one user session.

    Usage:
        ws = SessionWorkspace()
        path = ws.save_upload(uploaded_file.name, uploaded_file.getvalue())
        ... do work ...
        ws.cleanup()   # call this when the session/download is done

    Nothing here touches a database or any location outside the OS temp
    directory, and cleanup() removes it entirely.
    """

    def __init__(self):
        self.session_id = str(uuid.uuid4())
        self.dir = Path(tempfile.mkdtemp(prefix=f"sam_analyzer_{self.session_id}_"))

    def save_upload(self, filename: str, file_bytes: bytes) -> Path:
        safe_name = Path(filename).name  # strip any path components
        dest = self.dir / safe_name
        dest.write_bytes(file_bytes)
        return dest

    def cleanup(self):
        shutil.rmtree(self.dir, ignore_errors=True)

    def __enter__(self):
        return self

    def __exit__(self, exc_type, exc_val, exc_tb):
        self.cleanup()


# ---------------------------------------------------------------------------
# PDF parsing (native text + OCR fallback)
# ---------------------------------------------------------------------------

def _ocr_page(page: "pymupdf.Page") -> str:
    """
    Render a PDF page to an image and OCR it with Tesseract.

    Greyscale + autocontrast before OCR: contract scans are very often faxed,
    photocopied or photographed, and Tesseract's accuracy on low-contrast or
    yellowed pages improves markedly once the histogram is stretched. Cheap
    to do and it never hurts a clean scan.
    """
    zoom = OCR_RENDER_DPI / 72  # PDF base is 72 DPI
    matrix = pymupdf.Matrix(zoom, zoom)
    pix = page.get_pixmap(matrix=matrix)
    img = Image.open(io.BytesIO(pix.tobytes("png")))

    try:
        img = ImageOps.autocontrast(img.convert("L"))
    except Exception:
        pass  # preprocessing is an optimisation, never a hard requirement

    try:
        return pytesseract.image_to_string(img, config=OCR_CONFIG)
    except Exception:
        # A missing/misconfigured tesseract binary shouldn't kill the whole
        # document -- the native text (if any) is still worth returning.
        return ""


def _page_image_ratio(page: "pymupdf.Page") -> float:
    """Roughly what share of the page is covered by raster images."""
    try:
        page_area = abs(page.rect.width * page.rect.height)
        if page_area <= 0:
            return 0.0
        image_area = 0.0
        for info in page.get_image_info():
            bbox = info.get("bbox")
            if bbox:
                image_area += abs((bbox[2] - bbox[0]) * (bbox[3] - bbox[1]))
        return min(image_area / page_area, 1.0)
    except Exception:
        return 0.0


def _extract_tables(page: "pymupdf.Page") -> str:
    """
    Pull ruled/whitespace-aligned tables out of a page as pipe-separated rows.

    This matters more than it looks. Order forms, licence schedules and SKU
    tables are the single most valuable thing in a SAM contract, and plain
    text extraction flattens them column-by-column -- a part number ends up
    nowhere near its quantity, so the model has to guess which number belongs
    to which row, and quietly gets it wrong. Recovering the row structure and
    appending it in a labelled block gives the extractor something it can read
    a line item off unambiguously.
    """
    try:
        finder = page.find_tables()
    except Exception:
        return ""  # older PyMuPDF builds may not expose find_tables

    blocks = []
    for table in getattr(finder, "tables", []) or []:
        try:
            rows = table.extract()
        except Exception:
            continue
        rendered = [
            " | ".join((str(cell).strip() if cell is not None else "") for cell in row)
            for row in rows
            if any(cell is not None and str(cell).strip() for cell in row)
        ]
        if len(rendered) >= 2:  # a single row isn't a table worth restating
            blocks.append("\n".join(rendered))

    if not blocks:
        return ""
    return "\n\n[TABLE DETECTED ON THIS PAGE]\n" + "\n\n".join(blocks)


def parse_pdf(file_path: Path) -> ParsedDocument:
    """
    Page-by-page PDF extraction with three robustness measures aimed at the
    formats that actually turn up in a contract pack, as opposed to clean
    born-digital PDFs:

      1. Reading order -- sort=True orders text blocks top-to-bottom,
         left-to-right. Without it, two-column agreements and pages with
         side-margin definitions come out interleaved into nonsense.
      2. Tables -- extracted separately and appended, because flattened
         table text loses the row structure that makes a licence schedule
         readable (see _extract_tables).
      3. Hybrid OCR -- a page that has a little text AND a big image is
         treated as a scan with a digital header, not as a text page.
    """
    doc = ParsedDocument(filename=file_path.name, file_type="pdf")

    with pymupdf.open(file_path) as pdf:
        if pdf.needs_pass:
            # Try the empty password first: many contract PDFs are encrypted
            # only to restrict printing/copying, and open fine with no password.
            if not pdf.authenticate(""):
                raise ValueError(
                    "This PDF is password-protected. Remove the password and "
                    "upload it again."
                )

        for i, page in enumerate(pdf, start=1):
            try:
                native_text = page.get_text(sort=True).strip()
            except Exception:
                native_text = page.get_text().strip()

            table_text = _extract_tables(page)
            combined_native = (native_text + table_text).strip()

            if len(native_text) >= MIN_CHARS_FOR_NATIVE_TEXT:
                # Page has real text. But if there's little of it and the page
                # is mostly image, the substance is probably in the scan.
                if (len(native_text) < SPARSE_TEXT_THRESHOLD
                        and _page_image_ratio(page) > LARGE_IMAGE_AREA_RATIO):
                    ocr_text = _ocr_page(page).strip()
                    if len(ocr_text) > len(native_text):
                        merged = f"{combined_native}\n\n[OCR OF PAGE IMAGE]\n{ocr_text}"
                        doc.pages.append(PageResult(i, merged.strip(), "native+ocr"))
                        continue
                doc.pages.append(PageResult(i, combined_native, "native"))
            else:
                # Little or no extractable text -- scanned or blank page.
                ocr_text = _ocr_page(page).strip()
                if table_text:
                    ocr_text = (ocr_text + table_text).strip()
                doc.pages.append(PageResult(i, ocr_text, "ocr"))

    return doc


# ---------------------------------------------------------------------------
# Spreadsheet + plain-text parsing
# ---------------------------------------------------------------------------
# Order forms, licence schedules and renewal quotes routinely arrive as a
# spreadsheet rather than a signed PDF -- they're generated from a vendor's
# quoting system. Rejecting those forces someone to convert by hand, which is
# exactly the manual step this tool exists to remove.

def parse_spreadsheet(file_path: Path) -> ParsedDocument:
    """
    Flatten an .xlsx/.xlsm workbook to pipe-separated rows, one logical page
    per worksheet so a citation can still point at a named location.
    """
    from openpyxl import load_workbook

    def read(data_only: bool) -> list[tuple[str, list[str]]]:
        workbook = load_workbook(file_path, data_only=data_only, read_only=True)
        try:
            sheets = []
            for sheet in workbook.worksheets:
                rows = []
                for row in sheet.iter_rows(values_only=True):
                    cells = ["" if cell is None else str(cell).strip() for cell in row]
                    if any(cells):
                        rows.append(" | ".join(cells))
                if rows:
                    sheets.append((sheet.title, rows))
            return sheets
        finally:
            workbook.close()

    # data_only=True returns the cached values of formula cells rather than
    # the formula source -- a quote line reading "=B4*C4" helps nobody.
    sheets = read(data_only=True)

    # ...but the cache only exists if Excel itself last saved the file. A
    # workbook generated by a vendor's quoting system or by a script can have
    # every formula cell come back empty, which would silently drop the price
    # column of an order form. If the cached pass looks suspiciously thin,
    # fall back to the formula text -- ugly, but recoverable, and far better
    # than a blank column.
    cached_chars = sum(len(r) for _, rows in sheets for r in rows)
    if cached_chars < 40:
        fallback = read(data_only=False)
        if sum(len(r) for _, rows in fallback for r in rows) > cached_chars:
            sheets = fallback

    doc = ParsedDocument(filename=file_path.name, file_type="xlsx")
    for index, (title, rows) in enumerate(sheets, start=1):
        body = f"[WORKSHEET: {title}]\n" + "\n".join(rows)
        doc.pages.append(PageResult(index, body, "native"))

    if not doc.pages:
        raise ValueError("This spreadsheet appears to be empty.")
    return doc


def parse_delimited(file_path: Path) -> ParsedDocument:
    """CSV/TSV, with the delimiter sniffed rather than assumed."""
    raw = file_path.read_bytes().decode("utf-8", errors="replace")
    try:
        dialect = csv.Sniffer().sniff(raw[:4096])
        delimiter = dialect.delimiter
    except Exception:
        delimiter = ","

    rows = [
        " | ".join(cell.strip() for cell in row)
        for row in csv.reader(io.StringIO(raw), delimiter=delimiter)
        if any(cell.strip() for cell in row)
    ]
    if not rows:
        raise ValueError("This file appears to be empty.")

    doc = ParsedDocument(filename=file_path.name, file_type="csv")
    doc.pages.append(PageResult(1, "\n".join(rows), "native"))
    return doc


def parse_image(file_path: Path) -> ParsedDocument:
    """
    OCR a photo or scan of a contract page.

    Phone photos are the messiest input this tool takes: they arrive rotated,
    unevenly lit and low-contrast. Three cheap corrections handle most of it --
    EXIF rotation (a photo taken in portrait is stored sideways with an
    orientation tag Tesseract ignores), greyscale + autocontrast, and upscaling
    small images, since Tesseract needs roughly 300 DPI-equivalent resolution
    and a 1000px-wide photo of an A4 page is far short of that.
    """
    img = Image.open(file_path)

    try:
        img = ImageOps.exif_transpose(img)  # honour the camera's rotation tag
    except Exception:
        pass

    try:
        img = ImageOps.autocontrast(img.convert("L"))
    except Exception:
        pass

    # Upscale anything too small for reliable OCR. Below roughly 1600px across,
    # body text in a page photo falls under the ~20px character height
    # Tesseract wants and accuracy collapses.
    try:
        if img.width < 1600:
            scale = min(1600 / img.width, 3.0)
            img = img.resize(
                (int(img.width * scale), int(img.height * scale)),
                Image.LANCZOS,
            )
    except Exception:
        pass

    try:
        text = pytesseract.image_to_string(img, config=OCR_CONFIG).strip()
    except Exception as e:
        raise ValueError(
            "Couldn't run OCR on this image. Check that Tesseract is installed "
            f"({e})."
        ) from e

    if len(text) < MIN_CHARS_FOR_NATIVE_TEXT:
        raise ValueError(
            "No readable text could be found in this image. If it's a photo of "
            "a page, retake it straight-on in good light with the page filling "
            "the frame."
        )

    doc = ParsedDocument(filename=file_path.name, file_type="image")
    doc.pages.append(PageResult(1, text, "ocr"))
    return doc


def parse_plain_text(file_path: Path) -> ParsedDocument:
    text = file_path.read_bytes().decode("utf-8", errors="replace").strip()
    if not text:
        raise ValueError("This file appears to be empty.")
    doc = ParsedDocument(filename=file_path.name, file_type="txt")
    doc.pages.append(PageResult(1, text, "native"))
    return doc


# ---------------------------------------------------------------------------
# DOCX parsing
# ---------------------------------------------------------------------------

def parse_docx(file_path: Path) -> ParsedDocument:
    """
    DOCX has no native concept of "pages" (pagination is a rendering detail),
    so we treat the whole document as a single logical page, but we do walk
    paragraphs AND tables in document order so nothing gets silently dropped.
    """
    docx_doc = DocxDocument(file_path)

    chunks: list[str] = []
    for element in docx_doc.element.body:
        tag = element.tag.split("}")[-1]
        if tag == "p":
            para_text = "".join(node.text or "" for node in element.iter() if node.tag.endswith("}t"))
            if para_text.strip():
                chunks.append(para_text)
        elif tag == "tbl":
            chunks.append(_table_to_text(element))

    full_text = "\n".join(chunks).strip()

    doc = ParsedDocument(filename=file_path.name, file_type="docx")
    doc.pages.append(PageResult(page_number=1, text=full_text, source="native"))
    return doc


def _table_to_text(tbl_element) -> str:
    """Flatten a docx table XML element into pipe-separated rows."""
    rows_text = []
    for row in tbl_element.iter():
        if row.tag.endswith("}tr"):
            cells = [
                "".join(node.text or "" for node in cell.iter() if node.tag.endswith("}t"))
                for cell in row if cell.tag.endswith("}tc")
            ]
            rows_text.append(" | ".join(c.strip() for c in cells))
    return "\n".join(rows_text)


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------

def parse_document(file_path: Path) -> ParsedDocument:
    suffix = file_path.suffix.lower()
    if suffix == ".pdf":
        return parse_pdf(file_path)
    if suffix == ".docx":
        return parse_docx(file_path)
    if suffix in (".xlsx", ".xlsm"):
        return parse_spreadsheet(file_path)
    if suffix == ".csv":
        return parse_delimited(file_path)
    if suffix == ".txt":
        return parse_plain_text(file_path)
    if suffix in IMAGE_SUFFIXES:
        return parse_image(file_path)
    if suffix == ".doc":
        # Legacy binary .doc needs LibreOffice/antiword to convert; naming the
        # fix is more useful than a generic "unsupported" message.
        raise ValueError(
            "Legacy .doc files aren't supported. Open it in Word and save as "
            ".docx or print to PDF, then upload again."
        )
    raise ValueError(
        f"Unsupported file type: {suffix}. Supported: "
        + ", ".join(SUPPORTED_SUFFIXES)
    )
